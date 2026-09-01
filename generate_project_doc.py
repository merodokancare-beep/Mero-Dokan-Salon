import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Page Setup - Normal 1 inch margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

# Brand Colors
COLOR_PRIMARY = RGBColor(15, 42, 74)       # Deep Navy #0f2a4a
COLOR_SECONDARY = RGBColor(14, 116, 144)   # Deep Teal #0e7490
COLOR_TEXT = RGBColor(30, 41, 59)          # Dark Slate #1e293b
COLOR_MUTED = RGBColor(100, 116, 139)      # Muted Gray #64748b
COLOR_GOLD = RGBColor(180, 83, 9)          # Dark Gold #b45309
HEX_PRIMARY = "0f2a4a"
HEX_SECONDARY = "0e7490"
HEX_LIGHT_BG = "f8fafc"
HEX_GOLD_BG = "fffbeb"
HEX_BORDER = "cbd5e1"

def set_cell_shading(cell, color_hex):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_header_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY

def add_section_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY

def add_section_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY

def add_body(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Arial'
        r_bold.font.size = Pt(9.5)
        r_bold.font.bold = True
        r_bold.font.color.rgb = COLOR_TEXT
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(9.5)
    run.font.color.rgb = COLOR_TEXT

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Arial'
        r_bold.font.size = Pt(9.5)
        r_bold.font.bold = True
        r_bold.font.color.rgb = COLOR_TEXT
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(9.5)
    run.font.color.rgb = COLOR_TEXT

# ================= DOCUMENT START =================
add_header_title("MERODOKAN SALOON & SPA MANAGEMENT SYSTEM")
add_subtitle("Project Technical Documentation, Operational Workflows, Handover Sign-Off & AMC Agreement")

# Metadata Box Table
meta_table = doc.add_table(rows=1, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_table.autofit = False

left_cell = meta_table.cell(0, 0)
right_cell = meta_table.cell(0, 1)

set_cell_shading(left_cell, HEX_LIGHT_BG)
set_cell_shading(right_cell, HEX_LIGHT_BG)
set_cell_margins(left_cell, 120, 120, 150, 150)
set_cell_margins(right_cell, 120, 120, 150, 150)

p1 = left_cell.paragraphs[0]
p1.paragraph_format.space_after = Pt(2)
r1 = p1.add_run("Document Reference: ")
r1.font.bold = True
r1.font.size = Pt(9)
p1.add_run("DOC-SALOON-2026-v2.6\n").font.size = Pt(9)
r2 = p1.add_run("Software Product: ")
r2.font.bold = True
r2.font.size = Pt(9)
p1.add_run("MeroDokan Saloon & Spa POS (v2.6)\n").font.size = Pt(9)
r3 = p1.add_run("Architecture: ")
r3.font.bold = True
r3.font.size = Pt(9)
p1.add_run("100% Offline-First (MS SQL Server)").font.size = Pt(9)

p2 = right_cell.paragraphs[0]
p2.paragraph_format.space_after = Pt(2)
r4 = p2.add_run("Target Entity: ")
r4.font.bold = True
r4.font.size = Pt(9)
p2.add_run("Valued Salon / Spa Enterprise\n").font.size = Pt(9)
r5 = p2.add_run("Deployment Date: ")
r5.font.bold = True
r5.font.size = Pt(9)
p2.add_run("September 1, 2026\n").font.size = Pt(9)
r6 = p2.add_run("Status: ")
r6.font.bold = True
r6.font.size = Pt(9)
p2.add_run("Installed, Verified & Handed Over").font.size = Pt(9)

# 1. Executive Summary
add_section_h1("1. Executive Summary & Solution Architecture")
add_body("MeroDokan Saloon & Spa Management System is an enterprise desktop Point-of-Sale (POS), Customer Relationship Management (CRM), and resource planning suite engineered specifically for modern beauty salons, hair studios, unisex grooming lounges, cosmetic clinics, and luxury wellness spas. Built with a robust offline-first architecture, the platform guarantees zero counter downtime and instant sub-second billing response during peak hours.")

add_bullet(" Natively functions with zero internet connection requirement, eliminating cloud latency and downtime risks.", "100% Offline-First Reliability:")
add_bullet(" Simultaneously bills salon services and retail take-home cosmetic products on the same invoice.", "Dual-Item Hybrid Billing:")
add_bullet(" Automatically calculates service commission % and retail incentives per stylist row on each bill.", "Automated Stylist Commissions:")
add_bullet(" Prevents double-booking stylists with real-time clash verification and 1-click POS conversion.", "Visual Appointment Calendar:")
add_bullet(" Native ESC/POS printing for 58mm (2-inch) and 80mm (3-inch) receipts plus A4/A5 tax invoices.", "Multi-Format Thermal Receipts:")
add_bullet(" Automated encrypted daily database backups with 1-click restore protection.", "Data Privacy & Backups:")

# 2. Module Workflows
add_section_h1("2. Comprehensive Operational Workflows")

add_section_h2("2.1 Smart Express POS & Multi-Tender Checkout Workflow")
add_bullet(" Step 1 (Customer Lookup): Cashier enters client mobile number or searches by name; walk-in defaults to 'Walk-in Client'. Available loyalty points and outstanding dues display immediately.")
add_bullet(" Step 2 (Item Selection): Add services from touch categories (Hair, Skin, Spa) and scan retail products via barcode reader.")
add_bullet(" Step 3 (Stylist Allocation): Assign specific stylists to each individual service line item for accurate commission attribution.")
add_bullet(" Step 4 (Discounts & Tender): Apply item-level or overall discounts; choose payment method (Cash, UPI / QR, Card, Split Multi-Tender, or Customer Due/Credit).")
add_bullet(" Step 5 (Settlement & Print): Click 'Save & Print' to dispatch formatted thermal receipt, trigger cash drawer kick-out, and accrue loyalty points.")

add_section_h2("2.2 Appointment Booking & Lifecycle Management")
add_bullet(" Step 1 (Booking Creation): Select Customer, Date, Time Slot, Required Services, and Assigned Stylist.")
add_bullet(" Step 2 (Conflict Clash Detection): System dynamically checks stylist schedule; blocks overlapping bookings.")
add_bullet(" Step 3 (Visual Board Tracking): Stylist timeline columns display color-coded status (Scheduled, In-Progress, Completed, Cancelled).")
add_bullet(" Step 4 (1-Click Bill Conversion): Once service concludes, click 'Bill Now' to instantly populate the POS cart with zero re-entry.")

add_section_h2("2.3 Staff Management, Attendance & Commission Calculation")
add_bullet(" Step 1 (Roster Setup): Define staff roles (Senior Stylist, Junior Stylist, Beautician, Therapist) with base salaries and commission rates.")
add_bullet(" Step 2 (Daily Attendance): Staff record daily check-in / check-out timestamps and toggle availability status (Available, On Break, Busy).")
add_bullet(" Step 3 (Automated Calculation): Formula computes: Line Service Commission = (Price - Discount) * Commission %. Retail incentives tracked separately.")
add_bullet(" Step 4 (Payroll Reports): 1-Click export of monthly staff commission and productivity summaries.")

add_section_h2("2.4 Customer CRM, Dues Ledger & Loyalty Program")
add_bullet(" Captures client profiles, birthdays, anniversaries, and service notes.")
add_bullet(" Loyalty Engine: Accrues points based on bill amount; permits instant redemption as bill discounts.")
add_bullet(" Customer Due Ledger: Tracks unpaid credit balances with dedicated partial/full settlement receipts.")

add_section_h2("2.5 Inventory, Inward Purchase & Daily Stock Ledger")
add_bullet(" Inward purchase orders logged with supplier invoice number, batch details, expiry date, and landed cost.")
add_bullet(" Automatic separation between Retail Stock (for customer sales) and Back-Bar Stock (internal salon consumption).")
add_bullet(" Daily Stock Ledger Book provides real-time mathematical reconciliation: Closing = Opening + Inward - Sales - Internal Usage.")

add_section_h2("2.6 Day-End Cash Register Settlement (EOD Reconciliation)")
add_bullet(" Cashier inputs opening float at morning start.")
add_bullet(" System aggregates daily Cash Inflows + Digital Collections (UPI/Card) - Petty Cash Outflows.")
add_bullet(" Physical currency count is entered; system immediately computes and flags any cash shortage or surplus.")
add_bullet(" Generates official EOD Shift Handover slip for salon manager signing.")

# 3. UAT Verification Table
add_section_h1("3. User Acceptance Testing (UAT) & Deliverables Matrix")
uat_table = doc.add_table(rows=7, cols=3)
uat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
uat_table.autofit = False

headers = ["Module / Feature", "Verification & Deliverable Details", "UAT Status"]
for i, h in enumerate(headers):
    cell = uat_table.cell(0, i)
    set_cell_shading(cell, HEX_PRIMARY)
    set_cell_margins(cell, 80, 80, 100, 100)
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(255, 255, 255)

data = [
    ("Express POS Billing", "Dual service & retail items, dynamic stylist split, barcode scanner, discounts, split tenders", "PASSED (100%)"),
    ("Thermal Printing", "58mm & 80mm ESC/POS receipt generation with salon logo, tax breakdown, drawer kick", "PASSED (100%)"),
    ("Appointment Board", "Visual timeline calendar, clash conflict prevention, stylist board, 1-click billing", "PASSED (100%)"),
    ("Staff & Commissions", "Service % & retail incentive tracking, staff attendance check-in/out, payroll reports", "PASSED (100%)"),
    ("Stock & Ledger", "Purchase orders, supplier master, back-bar usage logs, low stock alerts, Daily Ledger", "PASSED (100%)"),
    ("EOD Settlement & CRM", "Day-end cash drawer reconciliation, loyalty points, customer credit ledger, 20+ reports", "PASSED (100%)")
]

for row_idx, (col1, col2, col3) in enumerate(data, start=1):
    c1 = uat_table.cell(row_idx, 0)
    c2 = uat_table.cell(row_idx, 1)
    c3 = uat_table.cell(row_idx, 2)
    
    if row_idx % 2 == 1:
        set_cell_shading(c1, HEX_LIGHT_BG)
        set_cell_shading(c2, HEX_LIGHT_BG)
        set_cell_shading(c3, HEX_LIGHT_BG)
        
    set_cell_margins(c1, 60, 60, 100, 100)
    set_cell_margins(c2, 60, 60, 100, 100)
    set_cell_margins(c3, 60, 60, 100, 100)
    
    c1.paragraphs[0].add_run(col1).font.size = Pt(8.5)
    c1.paragraphs[0].runs[0].font.bold = True
    c2.paragraphs[0].add_run(col2).font.size = Pt(8.5)
    
    r_stat = c3.paragraphs[0].add_run(col3)
    r_stat.font.size = Pt(8.5)
    r_stat.font.bold = True
    r_stat.font.color.rgb = RGBColor(22, 101, 52)

# 4. Project Sign-Off
add_section_h1("4. Official Project Completion & Handover Sign-Off Certificate")
add_body("This certificate confirms that the MeroDokan Saloon & Spa Management System (v2.6) has been successfully deployed, customized, tested, and commissioned. Full user training and operational documentation have been handed over to the Client.")

sig_table = doc.add_table(rows=2, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sig_table.autofit = False

s_th1 = sig_table.cell(0, 0)
s_th2 = sig_table.cell(0, 1)
set_cell_shading(s_th1, HEX_PRIMARY)
set_cell_shading(s_th2, HEX_PRIMARY)
set_cell_margins(s_th1, 80, 80, 120, 120)
set_cell_margins(s_th2, 80, 80, 120, 120)

s_th1.paragraphs[0].add_run("Delivered By (Service Provider)").font.color.rgb = RGBColor(255, 255, 255)
s_th1.paragraphs[0].runs[0].font.bold = True
s_th1.paragraphs[0].runs[0].font.size = Pt(9.5)

s_th2.paragraphs[0].add_run("Accepted & Confirmed By (Client / Salon)").font.color.rgb = RGBColor(255, 255, 255)
s_th2.paragraphs[0].runs[0].font.bold = True
s_th2.paragraphs[0].runs[0].font.size = Pt(9.5)

s_td1 = sig_table.cell(1, 0)
s_td2 = sig_table.cell(1, 1)
set_cell_margins(s_td1, 100, 100, 120, 120)
set_cell_margins(s_td2, 100, 100, 120, 120)

p_s1 = s_td1.paragraphs[0]
p_s1.add_run("Company: MeroDokan Software Solutions\nAuthorized Lead: Solutions Lead\nDate: September 1, 2026\n\nSignature: __________________________").font.size = Pt(9)

p_s2 = s_td2.paragraphs[0]
p_s2.add_run("Salon Name: _________________________\nAuthorized Person: ____________________\nDesignation: Owner / General Manager\n\nSignature & Seal: ____________________").font.size = Pt(9)

# 5. AMC Section
add_section_h1("5. Annual Maintenance Contract (AMC) Terms & SLA")
add_bullet(" Year 1 Warranty: 100% Complimentary (Included in deployment package).", "AMC Annual Fee:")
add_bullet(" Year 2 Onwards: Rs. 4,500 /- per year for Primary Node (Rs. 1,500 / yr per additional LAN terminal).", "Renewal Rate:")
add_bullet(" Unlimited remote desktop troubleshooting (AnyDesk/TeamViewer), minor software updates, SQL database health defragmentation, disaster recovery data restoration, and thermal printer recalibration.", "Scope of AMC:")
add_bullet(" P1 Critical (Counter Down): Response < 30 Mins | Resolution < 2 Hours.", "SLA Response Times:")
add_bullet(" P2 High (Feature Issue): Response < 2 Hours | Resolution < 6 Hours.", "SLA Response Times:")
add_bullet(" P3 Standard (General/Report Query): Response < 4 Hours | Resolution < 24 Hours.", "SLA Response Times:")

# 6. Terms & Conditions
add_section_h1("6. Standard Terms & Conditions")
add_bullet(" The software license is perpetual for the installed workstation. Client maintains 100% data ownership of their local database.", "Perpetual License:")
add_bullet(" AMC renewal fees are payable annually in advance. Non-renewal does not lock the offline POS system but shifts support to chargeable per-incident tickets.", "Payment Terms:")
add_bullet(" Client must maintain appropriate power backup (UPS) and safeguard automated daily database backups.", "Client Obligations:")

doc.save("MeroDokan_Saloon_Project_Documentation_SignOff_AMC.docx")
print("DOCX successfully generated: MeroDokan_Saloon_Project_Documentation_SignOff_AMC.docx")
