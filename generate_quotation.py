import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def build_valid_docx():
    doc = docx.Document()
    
    # ---------------- PAGE MARGINS ----------------
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11.0)
        
        # Header / Footer
        footer = sec.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_f = f_p.add_run("MeroDokan Saloon & Spa POS System | Commercial Quotation & Proposal")
        r_f.font.name = "Arial"
        r_f.font.size = Pt(8)
        r_f.font.color.rgb = RGBColor(120, 120, 120)

    # Color Palette
    NAVY = RGBColor(15, 42, 74)        # #0F2A4A Primary Dark
    TEAL = RGBColor(14, 116, 144)      # #0E7490 Accent Teal
    CHARCOAL = RGBColor(45, 55, 72)    # #2D3748 Body
    DARK_BLUE = RGBColor(0, 70, 140)

    def add_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(20)
        r.font.bold = True
        r.font.color.rgb = NAVY
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(11)
        r.font.color.rgb = TEAL
        return p

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = NAVY
        return p

    def add_body(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.font.name = "Arial"
            r_b.font.size = Pt(9.5)
            r_b.font.bold = True
            r_b.font.color.rgb = NAVY
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        r.font.color.rgb = CHARCOAL
        return p

    def add_bullet(bold_txt, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        if bold_txt:
            r1 = p.add_run(bold_txt + ": ")
            r1.font.name = "Arial"
            r1.font.size = Pt(9.5)
            r1.font.bold = True
            r1.font.color.rgb = NAVY
        r2 = p.add_run(text)
        r2.font.name = "Arial"
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = CHARCOAL
        return p

    def set_cell_shading(cell, color_hex):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    # ---------------- COVER / HEADER ----------------
    add_title("COMMERCIAL PROPOSAL & QUOTATION")
    add_subtitle("MeroDokan Saloon & Spa Management System (Enterprise POS & CRM Suite)")

    # Metadata Box
    meta_tbl = doc.add_table(rows=2, cols=2)
    meta_tbl.style = 'Table Grid'
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_tbl.autofit = False

    meta_tbl.rows[0].cells[0].width = Inches(3.5)
    meta_tbl.rows[0].cells[1].width = Inches(3.5)
    meta_tbl.rows[1].cells[0].width = Inches(3.5)
    meta_tbl.rows[1].cells[1].width = Inches(3.5)

    for row in meta_tbl.rows:
        for c in row.cells:
            set_cell_shading(c, "F1F5F9")

    # Cell 0,0
    p = meta_tbl.rows[0].cells[0].paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Quotation Ref: "); r.bold = True; r.font.size = Pt(9); r.font.name = "Arial"
    r2 = p.add_run("QTN-2026-SALOON-883\n"); r2.font.size = Pt(9); r2.font.name = "Arial"
    r3 = p.add_run("Proposal Date: "); r3.bold = True; r3.font.size = Pt(9); r3.font.name = "Arial"
    r4 = p.add_run("23 August 2026"); r4.font.size = Pt(9); r4.font.name = "Arial"

    # Cell 0,1
    p = meta_tbl.rows[0].cells[1].paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Prepared For: "); r.bold = True; r.font.size = Pt(9); r.font.name = "Arial"
    r2 = p.add_run("Valued Salon / Spa Owner & Management\n"); r2.font.size = Pt(9); r2.font.name = "Arial"
    r3 = p.add_run("Industry: "); r3.bold = True; r3.font.size = Pt(9); r3.font.name = "Arial"
    r4 = p.add_run("Unisex Salon, Hair Studio, Beauty Spa & Parlour"); r4.font.size = Pt(9); r4.font.name = "Arial"

    # Cell 1,0
    p = meta_tbl.rows[1].cells[0].paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Software Solution: "); r.bold = True; r.font.size = Pt(9); r.font.name = "Arial"
    r2 = p.add_run("MeroDokan Saloon & Spa POS (v2.6)\n"); r2.font.size = Pt(9); r2.font.name = "Arial"
    r3 = p.add_run("Architecture: "); r3.bold = True; r3.font.size = Pt(9); r3.font.name = "Arial"
    r4 = p.add_run("High-Speed Desktop (.NET + SQL Database)"); r4.font.size = Pt(9); r4.font.name = "Arial"

    # Cell 1,1
    p = meta_tbl.rows[1].cells[1].paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Proposal Validity: "); r.bold = True; r.font.size = Pt(9); r.font.name = "Arial"
    r2 = p.add_run("30 Days from date of issuance\n"); r2.font.size = Pt(9); r2.font.name = "Arial"
    r3 = p.add_run("Turnkey Delivery: "); r3.bold = True; r3.font.size = Pt(9); r3.font.name = "Arial"
    r4 = p.add_run("1 to 2 Business Days"); r4.font.size = Pt(9); r4.font.name = "Arial"

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------- 1. EXECUTIVE SUMMARY ----------------
    add_heading_1("1. Executive Summary & Solution Overview")
    add_body(
        "MeroDokan Saloon & Spa Management System is an enterprise-grade, desktop-based Point of Sale (POS) and "
        "business automation platform engineered exclusively for modern beauty salons, hair studios, luxury spas, and cosmetic parlours. "
        "The system replaces fragmented manual registers with a unified, high-speed software that streamlines daily billing, "
        "tracks stylist service commissions automatically, manages appointments to prevent double-booking, audits product stock, and generates comprehensive statutory tax and business analytics reports."
    )

    # Highlights Box
    hl_tbl = doc.add_table(rows=1, cols=1)
    hl_tbl.style = 'Table Grid'
    hl_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hl_tbl.autofit = False
    hl_tbl.rows[0].cells[0].width = Inches(7.0)
    set_cell_shading(hl_tbl.rows[0].cells[0], "E0F2FE")
    p_hl = hl_tbl.rows[0].cells[0].paragraphs[0]
    p_hl.paragraph_format.space_after = Pt(2)
    r_h = p_hl.add_run("★ Key Business Advantages:\n")
    r_h.bold = True; r_h.font.size = Pt(9.5); r_h.font.color.rgb = DARK_BLUE; r_h.font.name = "Arial"
    r_b = p_hl.add_run(
        "• 100% Offline-First Architecture: Zero counter lag, works without internet connectivity.\n"
        "• Automated Stylist Commissions: Tracks service & retail sales commissions per staff member.\n"
        "• Visual Appointment Calendar: Schedule clients, assign stylists, and convert bookings to bills in 1 click.\n"
        "• Multi-Mode Billing & Thermal Receipts: Fast checkout with 58mm/80mm thermal receipts and A4/A5 invoices.\n"
        "• Inventory & Consumption Tracking: Separate retail sales stock and in-salon back-bar consumption.\n"
        "• Full Data Security & Local Backups: Your customer data and financial numbers stay 100% private."
    )
    r_b.font.size = Pt(9); r_b.font.color.rgb = CHARCOAL; r_b.font.name = "Arial"

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------- 2. DETAILED MODULE BREAKDOWN ----------------
    add_heading_1("2. Comprehensive Software Modules & Functionality")
    add_body("The complete software suite is bundled into a single turnkey license with all modules unlocked:")

    modules = [
        ("1. Smart POS & Express Billing", [
            "Lightning-fast billing with touchscreen and keyboard shortcut support.",
            "Dual-item billing in a single invoice: Salon Services (Haircut, Facial, Spa) + Retail Products (Shampoos, Serums).",
            "Stylist/Staff assignment per individual service row with automatic commission tracking.",
            "Multiple payment modes: Cash, Card, UPI / QR Codes, Digital Wallets, Split Payments, and Customer Due/Credit.",
            "Fast 58mm (2-inch) and 80mm (3-inch) ESC/POS Thermal Receipt printing + Standard A4/A5 Invoicing.",
            "Bill Hold & Resume ('Park Bill') to handle multiple clients in parallel at busy counters.",
            "Item-wise discount and overall bill discount (Percentage % or Flat Amount).",
            "Barcode scanner integration for instant retail product lookup."
        ]),
        ("2. Appointment Scheduling & Visual Calendar", [
            "Interactive daily and weekly booking calendar view.",
            "Time-slot scheduling with real-time clash detection to prevent double-booking stylists.",
            "Booking lifecycle tracking: Scheduled, In-Progress, Completed, Cancelled, No-Show.",
            "1-Click direct conversion from completed appointment into active POS bill.",
            "Customer service preferences and appointment notes history."
        ]),
        ("3. Service Catalog & Package Management", [
            "Complete service catalog with categorised services (Hair, Skin, Spa, Bridal, Beard, Nails).",
            "Custom service duration and pricing tier management.",
            "Default commission % or fixed commission amount defined per service.",
            "HSN/SAC code tax mapping for statutory compliance."
        ]),
        ("4. Inventory, Stock Ledger & Retail Management", [
            "Real-time product stock tracking for both retail sales and in-salon back-bar usage.",
            "Purchase order entry with supplier invoices, cost prices, batch details, and GST/VAT.",
            "Low-stock threshold alerts with instant warnings on dashboard.",
            "Comprehensive Stock Ledger: Chronological audit of inward purchases, outward sales, usage, and returns.",
            "Built-in Barcode Generator and label printing module for non-barcoded salon products."
        ]),
        ("5. Customer CRM & Loyalty Reward Program", [
            "Centralized customer directory (Mobile, Name, Gender, Birthday, Anniversary, Address).",
            "Automated Loyalty Points engine: Accumulate points on visits, redeem points during checkout.",
            "Customer visit history, total lifetime spend, and preferred staff members.",
            "Customer credit ledger for tracking dues and balance settlement.",
            "Birthday & Anniversary alerts for targeted promotions and customer retention."
        ]),
        ("6. Staff & Stylist Commission Management", [
            "Staff master database with role definitions (Stylists, Beauticians, Therapists, Cashiers).",
            "Automated commission calculation for service revenue and retail product sales.",
            "Comprehensive staff performance reports: Services performed, tips, revenue generated, and commission payout.",
            "Staff attendance and daily check-in / check-out availability tracking."
        ]),
        ("7. Day-End Cash Register & Daily Settlement (EOD)", [
            "End-of-Day (EOD) register closing and cash drawer reconciliation.",
            "Automated formula: Opening Cash + Cash Sales + Digital Collections - Cash Outflow = Expected Balance.",
            "Discrepancy detection to identify cash shortages or overages immediately.",
            "Printable Daily Settlement summary sheet for salon owners and accountants."
        ]),
        ("8. Tax, GST / VAT & HSN/SAC Compliance", [
            "Built-in HSN/SAC code master with standard GST / VAT tax rate slabs.",
            "Automatic CGST / SGST / IGST tax split on bills.",
            "Compliant tax invoice generation and audit-ready tax reports exportable to Excel."
        ]),
        ("9. Business Intelligence & Reporting Suite (20+ Reports)", [
            "Sales Analysis: Daily, Monthly, Date-range, Category-wise, Service-wise, and Payment-mode breakdown.",
            "Staff Commission Reports: Detailed per-stylist commission breakdown for payroll.",
            "Inventory Valuation & Stock Movement reports.",
            "Customer Retention & Spending reports.",
            "Profit & Loss / Gross Margin estimation.",
            "1-Click Export to Excel, CSV, PDF, and direct printing."
        ]),
        ("10. Security, Multi-User & Backup Management", [
            "Role-Based Access Control (Admin, Manager, Cashier, Receptionist, Staff).",
            "Secure encrypted password authentication and active session management.",
            "Automated daily database backups and 1-click manual backup/restore.",
            "Hardware-locked machine licensing protection with custom salon branding."
        ])
    ]

    for mod_title, mod_features in modules:
        p_m = doc.add_paragraph()
        p_m.paragraph_format.space_before = Pt(6)
        p_m.paragraph_format.space_after = Pt(2)
        r_mt = p_m.add_run("■ " + mod_title)
        r_mt.font.name = "Arial"
        r_mt.font.size = Pt(10.5)
        r_mt.font.bold = True
        r_mt.font.color.rgb = DARK_BLUE
        
        for feat in mod_features:
            add_bullet("", feat)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------- 3. TECHNICAL SPECIFICATIONS ----------------
    add_heading_1("3. System Requirements & Hardware Compatibility")
    
    req_tbl = doc.add_table(rows=6, cols=3)
    req_tbl.style = 'Table Grid'
    req_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    req_tbl.autofit = False

    req_headers = ["Hardware Component", "Minimum Requirement", "Recommended Specification"]
    widths = [Inches(2.2), Inches(2.4), Inches(2.4)]
    
    for c_i, h_txt in enumerate(req_headers):
        cell = req_tbl.rows[0].cells[c_i]
        cell.width = widths[c_i]
        set_cell_shading(cell, "0F2A4A")
        p = cell.paragraphs[0]
        r = p.add_run(h_txt)
        r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(255, 255, 255); r.font.name = "Arial"

    req_rows = [
        ("Operating System", "Windows 10 (64-bit)", "Windows 11 / Windows 10 Pro (64-bit)"),
        ("Processor (CPU)", "Intel Core i3 (4th Gen) / AMD Ryzen 3", "Intel Core i5 (8th Gen+) / AMD Ryzen 5"),
        ("System Memory (RAM)", "4 GB RAM", "8 GB / 16 GB DDR4 RAM"),
        ("Disk Storage", "500 MB Free Space (HDD)", "Fast SSD (Solid State Drive)"),
        ("Peripherals / Printers", "58mm / 80mm USB Thermal Printer", "80mm ESC/POS Thermal Printer + Barcode Scanner")
    ]

    for r_idx, row_data in enumerate(req_rows):
        row = req_tbl.rows[r_idx + 1]
        bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            cell.width = widths[c_i]
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(8.5); r.font.name = "Arial"
            if c_i == 0:
                r.font.bold = True
                r.font.color.rgb = NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------- 4. SCOPE OF IMPLEMENTATION ----------------
    add_heading_1("4. Scope of Implementation & Turnkey Deliverables")
    add_bullet("Software Deployment", "Complete automated installation and database configuration on counter workstation(s).")
    add_bullet("Branding & Customization", "Salon logo upload, company details, GSTIN/PAN setup, and custom thermal receipt footer notes.")
    add_bullet("Master Data Setup", "Assistance in populating service categories, pricing, retail items, opening inventory, and staff master.")
    add_bullet("Hardware Calibration", "Testing and driver configuration for thermal receipt printers (58mm/80mm) and barcode scanners.")
    add_bullet("Staff Onboarding & Training", "Hands-on training session for salon managers, front-desk receptionists, and cashiers.")
    add_bullet("12 Months Support & Warranty", "Dedicated technical support, software bug resolution, database maintenance, and remote assistance.")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------- 5. COMMERCIAL QUOTATION TABLE ----------------
    add_heading_1("5. Commercial Quotation & Investment Schedule")
    add_body("Below is the itemized quotation for the software package, setup, and support services:")

    pr_tbl = doc.add_table(rows=6, cols=4)
    pr_tbl.style = 'Table Grid'
    pr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    pr_tbl.autofit = False

    p_headers = ["Item / Description", "License Type", "Deliverables", "Price (INR / NPR)"]
    p_widths = [Inches(2.5), Inches(1.3), Inches(2.0), Inches(1.2)]
    
    for c_i, h_txt in enumerate(p_headers):
        cell = pr_tbl.rows[0].cells[c_i]
        cell.width = p_widths[c_i]
        set_cell_shading(cell, "0F2A4A")
        p = cell.paragraphs[0]
        r = p.add_run(h_txt)
        r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(255, 255, 255); r.font.name = "Arial"
        if c_i == 3:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    pricing_data = [
        ("MeroDokan Saloon & Spa POS (Enterprise Edition)\n• Full software license with all 10 core modules unlocked",
         "Lifetime License\n(Perpetual)",
         "Full Software + DB Engine + 1 Year Free Version Updates",
         "Rs. 25,000 /-"),
        ("Turnkey Setup & Custom Branding\n• Installation, thermal printer setup, logo branding & initial data entry",
         "One-Time Setup",
         "Onsite / Remote Installation + Hardware Configuration",
         "INCLUDED\n(FREE)"),
        ("Staff Training & Cashier Onboarding\n• Interactive training for salon manager, receptionists & staff",
         "Included",
         "Up to 2 Training Sessions + Digital User Manual",
         "INCLUDED\n(FREE)"),
        ("Annual Maintenance & Priority Support (AMC)\n• Dedicated remote AnyDesk support & database health check",
         "Annual (Optional)",
         "1st Year 100% FREE; Renewal applicable from 2nd Year",
         "Rs. 4,500 / yr\n(Optional)"),
        ("Additional Terminal License (Secondary Counter / Back Office)\n• For multi-counter salon setups connecting to main database",
         "Per Node",
         "Secondary Terminal License connected to LAN database",
         "Rs. 6,000 / node\n(Optional)")
    ]

    for r_idx, row_data in enumerate(pricing_data):
        row = pr_tbl.rows[r_idx + 1]
        bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            cell.width = p_widths[c_i]
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(8.5); r.font.name = "Arial"
            if c_i == 0:
                r.font.bold = True; r.font.color.rgb = NAVY
            elif c_i == 3:
                r.font.bold = True; r.font.color.rgb = TEAL
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Summary Callout
    tot_tbl = doc.add_table(rows=1, cols=1)
    tot_tbl.style = 'Table Grid'
    tot_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tot_tbl.autofit = False
    tot_tbl.rows[0].cells[0].width = Inches(7.0)
    set_cell_shading(tot_tbl.rows[0].cells[0], "FEF3C7")
    p_tot = tot_tbl.rows[0].cells[0].paragraphs[0]
    p_tot.paragraph_format.space_after = Pt(2)
    r_tot_h = p_tot.add_run("★ TOTAL TURNKEY PACKAGE INVESTMENT: Rs. 25,000 /- (All-Inclusive)\n")
    r_tot_h.bold = True; r_tot_h.font.size = Pt(10); r_tot_h.font.color.rgb = RGBColor(146, 64, 14); r_tot_h.font.name = "Arial"
    r_tot_b = p_tot.add_run(
        "Includes Lifetime Perpetual License for Primary Terminal + Complete Setup & Custom Branding + Staff Training + 1 Year Dedicated Technical Support."
    )
    r_tot_b.font.size = Pt(9); r_tot_b.font.color.rgb = CHARCOAL; r_tot_b.font.name = "Arial"

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------- 6. TERMS & CONDITIONS ----------------
    add_heading_1("6. Commercial Terms & Payment Milestones")
    add_bullet("Payment Schedule", "50% advance upon Work Order confirmation; remaining 50% upon successful software deployment and staff training.")
    add_bullet("Statutory Taxes", "Applicable statutory government taxes (e.g. GST/VAT) extra as per prevailing government laws if tax invoice is required.")
    add_bullet("Delivery Timeline", "Software installation and hardware calibration completed within 1–2 business days of order confirmation.")
    add_bullet("Hardware Obligation", "Client to provide the PC workstation, thermal printer, and barcode scanner as per technical specifications.")
    add_bullet("Warranty & Support", "12 Months free remote assistance via AnyDesk, version updates, and database maintenance advisory.")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 7. ACCEPTANCE & SIGN-OFF ----------------
    add_heading_1("7. Acceptance & Authorization Sign-Off")
    add_body("Please sign and return this page to confirm your acceptance of the proposal and initiate deployment:")

    sig_tbl = doc.add_table(rows=4, cols=2)
    sig_tbl.style = 'Table Grid'
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_tbl.autofit = False

    sig_tbl.rows[0].cells[0].width = Inches(3.5)
    sig_tbl.rows[0].cells[1].width = Inches(3.5)
    sig_tbl.rows[1].cells[0].width = Inches(3.5)
    sig_tbl.rows[1].cells[1].width = Inches(3.5)
    sig_tbl.rows[2].cells[0].width = Inches(3.5)
    sig_tbl.rows[2].cells[1].width = Inches(3.5)
    sig_tbl.rows[3].cells[0].width = Inches(3.5)
    sig_tbl.rows[3].cells[1].width = Inches(3.5)

    # Header
    set_cell_shading(sig_tbl.rows[0].cells[0], "0F2A4A")
    set_cell_shading(sig_tbl.rows[0].cells[1], "0F2A4A")
    r1 = sig_tbl.rows[0].cells[0].paragraphs[0].add_run("Authorized By (Service Provider)")
    r1.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = RGBColor(255, 255, 255); r1.font.name = "Arial"
    r2 = sig_tbl.rows[0].cells[1].paragraphs[0].add_run("Accepted By (Client / Salon Management)")
    r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(255, 255, 255); r2.font.name = "Arial"

    # Data
    fields = [
        ("Representative: Software Solutions / MeroDokan Team", "Client / Business Name: _______________________"),
        ("Designation: Solutions & Technical Lead", "Authorized Person: ___________________________"),
        ("Signature & Seal: ____________________________", "Signature & Seal: ____________________________")
    ]

    for r_idx, (p_text, c_text) in enumerate(fields):
        row = sig_tbl.rows[r_idx + 1]
        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r = p1.add_run(p_text); r.font.size = Pt(8.5); r.font.name = "Arial"
        
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(2)
        r = p2.add_run(c_text); r.font.size = Pt(8.5); r.font.name = "Arial"

    # Save Document
    doc.save("MeroDokan_Saloon_POS_Commercial_Quotation.docx")
    print("MeroDokan_Saloon_POS_Commercial_Quotation.docx generated successfully!")

if __name__ == "__main__":
    build_valid_docx()
